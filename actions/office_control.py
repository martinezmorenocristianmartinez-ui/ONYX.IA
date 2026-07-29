"""office_control.py — Advanced Word, Excel, and Notepad automation for ONYX."""
import os
import json
import subprocess
import time
from pathlib import Path
from datetime import datetime
from core.file_utils import force_save_file


def _ensure_com():
    try:
        import pythoncom
        pythoncom.CoInitialize()
    except Exception:
        pass


def open_blank_word():
    """Abre Microsoft Word con un documento en blanco usando COM."""
    try:
        _ensure_com()
        import win32com.client
        for _ in range(3):
            try:
                word = win32com.client.GetActiveObject("Word.Application")
                word.Visible = True
                word.Documents.Add()
                return "He abierto Microsoft Word con un documento en blanco."
            except Exception:
                import time
                time.sleep(1.0)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True
        word.Documents.Add()
        return "He abierto Microsoft Word con un documento en blanco."
    except Exception as e:
        os.system('start winword.exe /t')
        return f"He abierto Word (fallback)."


def open_blank_excel():
    """Abre Microsoft Excel con una hoja en blanco usando COM."""
    try:
        _ensure_com()
        import win32com.client
        for _ in range(3):
            try:
                excel = win32com.client.GetActiveObject("Excel.Application")
                excel.Visible = True
                excel.Workbooks.Add()
                return "He abierto Microsoft Excel con una hoja en blanco."
            except Exception:
                import time
                time.sleep(1.0)
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = True
        excel.Workbooks.Add()
        return "He abierto Microsoft Excel con una hoja en blanco."
    except Exception as e:
        os.system('start excel.exe')
        return f"He abierto Excel (fallback)."


def open_blank_powerpoint():
    """Abre Microsoft PowerPoint con una presentación en blanco."""
    try:
        os.system('start powerpnt.exe')
        return "He abierto Microsoft PowerPoint con una presentación en blanco."
    except Exception as e:
        return f"Error al abrir PowerPoint: {str(e)}"


def office_control(parameters: dict, player=None) -> str:
    """
    Handle Word, Excel, and Notepad operations.
    """
    action = parameters.get("action", "").lower()
    file_path = parameters.get("file_path", "")
    content = parameters.get("content", {})
    text = parameters.get("text", content.get("body", ""))

    # Acciones de abrir en blanco
    if action in ["open_blank_word", "blank_word", "abrir_word_blanco", "abrir_hoja_blanca_word", "word_blank"]:
        return open_blank_word()
    elif action in ["open_blank_excel", "blank_excel", "abrir_excel_blanco", "abrir_hoja_blanca_excel", "excel_blank"]:
        return open_blank_excel()
    elif action in ["open_blank_powerpoint", "blank_powerpoint", "abrir_powerpoint_blanco", "abrir_presentacion_blanca", "powerpoint_blank"]:
        return open_blank_powerpoint()

    if not file_path and action not in ["open_word", "open_excel", "open_notepad"]:
        desktop = Path(os.path.join(os.environ["USERPROFILE"], "Desktop"))
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if "doc" in action:
            file_path = str(desktop / f"Documento_{timestamp}.docx")
        elif "sheet" in action or "excel" in action:
            file_path = str(desktop / f"Planilla_{timestamp}.xlsx")
        elif "notepad" in action or "txt" in action:
            file_path = str(desktop / f"Nota_{timestamp}.txt")

    try:
        # Word operations
        if action in ["create_doc", "create_word", "create_documento"]:
            from docx import Document
            def _save_word(path):
                doc = Document()
                doc.add_heading(content.get("title", "Documento de ONYX"), 0)
                body = content.get("body", text)
                if isinstance(body, list):
                    for p in body:
                        doc.add_paragraph(p)
                else:
                    doc.add_paragraph(body)
                doc.save(path)

            result_msg = force_save_file(file_path, _save_word)
            if player:
                player.write_log(f"📝 {result_msg}")
            return f"Hecho, Señor Cristian. {result_msg}"

        elif action in ["edit_doc", "edit_word", "editar_documento"]:
            from docx import Document
            doc = Document(file_path)
            body = content.get("body", text)
            if isinstance(body, list):
                for p in body:
                    doc.add_paragraph(p)
            else:
                doc.add_paragraph(body)
            doc.save(file_path)
            return f"He actualizado el documento Word, Señor Cristian."

        elif action in ["open_word", "abrir_word", "abrir_documento_word"]:
            os.system('start winword.exe')
            if player:
                player.write_log("🚀 Abriendo Microsoft Word...")
            return "He abierto Word, Señor Cristian."

        # Excel operations
        elif action in ["create_sheet", "create_excel", "create_planilla", "crear_excel"]:
            from openpyxl import Workbook
            def _save_excel(path):
                wb = Workbook()
                ws = wb.active
                ws.title = content.get("sheet_name", "ONYX Data")
                data = content.get("data", [])
                for row in data:
                    ws.append(row)
                wb.save(path)

            result_msg = force_save_file(file_path, _save_excel)
            if player:
                player.write_log(f"📊 {result_msg}")
            return f"Hecho, Señor Cristian. {result_msg}"

        elif action in ["edit_sheet", "edit_excel", "editar_planilla"]:
            from openpyxl import load_workbook
            wb = load_workbook(file_path)
            ws = wb.active
            data = content.get("data", [])
            for row in data:
                ws.append(row)
            wb.save(file_path)
            return f"He actualizado el archivo Excel, Señor Cristian."

        elif action in ["open_excel", "abrir_excel"]:
            os.system('start excel.exe')
            if player:
                player.write_log("🚀 Abriendo Microsoft Excel...")
            return "He abierto Excel, Señor Cristian."

        # Notepad operations
        elif action in ["create_notepad", "create_txt", "crear_nota", "crear_texto"]:
            def _save_txt(path):
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(text)

            result_msg = force_save_file(file_path, _save_txt)
            if player:
                player.write_log(f"📄 {result_msg}")
            return f"Hecho, Señor Cristian. {result_msg}"

        elif action in ["edit_notepad", "edit_txt", "editar_nota"]:
            with open(file_path, 'a', encoding='utf-8') as f:
                f.write(f"\n{text}")
            return f"He agregado texto al bloc de notas, Señor Cristian."

        elif action in ["open_notepad", "abrir_bloc_de_notas", "abrir_notepad"]:
            if file_path:
                os.system(f'start notepad.exe "{file_path}"')
                if player:
                    player.write_log(f"📄 Abriendo Notepad con {Path(file_path).name}...")
                return f"He abierto el bloc de notas con el archivo, Señor Cristian."
            else:
                os.system('start notepad.exe')
                if player:
                    player.write_log("🚀 Abriendo Notepad...")
                return "He abierto el bloc de notas, Señor Cristian."

        return f"Acción de Office '{action}' no reconocida."

    except Exception as e:
        return f"Error en la operación de Office, Señor Cristian: {str(e)}"
