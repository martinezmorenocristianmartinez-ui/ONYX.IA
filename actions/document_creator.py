import os
from pathlib import Path
from datetime import datetime
from core.file_utils import force_save_file, get_desktop_path

def document_creator(parameters: dict, player=None) -> str:
    """
    Crea documentos de texto, Word o Excel con guardado forzado (Claude Code style).
    """
    action = parameters.get("action", "").lower()
    title = parameters.get("title", "Documento_Sin_Titulo")
    content = parameters.get("content", "")
    sheets = parameters.get("sheets", [])
    
    desktop_path = get_desktop_path()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ']).rstrip().replace(" ", "_")
    if not safe_title: safe_title = "Documento"

    try:
        if action == "word" or action == "google_doc":
            from docx import Document
            def _save_word(path):
                doc = Document()
                doc.add_heading(title, 0)
                lines = content.split('\n')
                table_rows = []
                in_table = False
                for line in lines:
                    line_stripped = line.strip()
                    if not line_stripped: continue
                    if line_stripped.startswith('|') and line_stripped.endswith('|'):
                        cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
                        if not any('---' in c for c in cells):
                            table_rows.append(cells)
                            in_table = True
                            continue
                        else:
                            continue
                    if in_table and table_rows:
                        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                        table.style = 'Table Grid'
                        for i, row_data in enumerate(table_rows):
                            for j, cell_data in enumerate(row_data):
                                table.rows[i].cells[j].text = cell_data
                        table_rows = []
                        in_table = False
                    if line_stripped.startswith('## '): doc.add_heading(line_stripped[3:], level=2)
                    elif line_stripped.startswith('# '): doc.add_heading(line_stripped[2:], level=1)
                    elif line_stripped.startswith('- '): doc.add_paragraph(line_stripped[2:], style='List Bullet')
                    else: doc.add_paragraph(line_stripped)
                if in_table and table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_rows):
                        for j, cell_data in enumerate(row_data):
                            table.rows[i].cells[j].text = cell_data
                doc.save(path)
            
            file_path = desktop_path / f"{safe_title}_{timestamp}.docx"
            result_msg = force_save_file(str(file_path), _save_word)
            return f"Hecho, Señor Cristian. {result_msg}"

        elif action == "excel" or action == "google_sheet":
            from openpyxl import Workbook
            def _save_excel(path):
                wb = Workbook()
                wb.remove(wb.active)
                if not sheets: return
                for sheet_data in sheets:
                    ws = wb.create_sheet(title=sheet_data.get("name", "Hoja")[:31])
                    if sheet_data.get("headers"): ws.append(sheet_data.get("headers"))
                    for row in sheet_data.get("rows", []): ws.append(row)
                wb.save(path)
                
            file_path = desktop_path / f"{safe_title}_{timestamp}.xlsx"
            result_msg = force_save_file(str(file_path), _save_excel)
            return f"Hecho, Señor Cristian. {result_msg}"

        elif action == "text":
            def _save_text(path):
                with open(path, "w", encoding="utf-8") as f:
                    f.write(f"{title}\n\n{content}")
            
            file_path = desktop_path / f"{safe_title}_{timestamp}.txt"
            result_msg = force_save_file(str(file_path), _save_text)
            return f"Hecho, Señor Cristian. {result_msg}"
            
        else:
            return f"Acción '{action}' desconocida. Use 'word', 'excel' o 'text'."
            
    except Exception as e:
        return f"Error al crear el documento: {str(e)}"
